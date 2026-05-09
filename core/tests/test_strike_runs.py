"""Tests for the unified strike model [D-060].

Three layers:

1. **Model** — ``ContentBlock.live_text``, ``row_all_struck``,
   ``header_all_struck``, JSON roundtrip preserves runs.
2. **DOCX extractor** — runs populated for paragraphs / table cells;
   per-cell + per-row strike marked, never dropped; whole-table strike
   on font_info.strikethrough; partial strike leaves the flag False.
3. **Parser** — partial-struck paragraphs become live_text in the parsed
   tree; struck rows are dropped from chunked output; req_ids in struck
   spans land in struck_req_ids and are excluded from the tree.
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from core.src.extraction.docx_extractor import DOCXExtractor
from core.src.models.document import (
    BlockType,
    ContentBlock,
    DocumentIR,
    FontInfo,
    Position,
    TextRun,
)


# ---------------------------------------------------------------------------
# Model layer
# ---------------------------------------------------------------------------

class TestLiveText:
    def test_partial_strike_returns_unstruck_concat(self):
        b = ContentBlock(
            type=BlockType.PARAGRAPH,
            position=Position(page=1, index=0),
            text="The UE shall not retry within 30s",
            runs=[
                TextRun(text="The UE shall ", struck=False),
                TextRun(text="not ", struck=True),
                TextRun(text="retry within 30s", struck=False),
            ],
        )
        assert b.live_text() == "The UE shall retry within 30s"

    def test_full_strike_returns_empty(self):
        b = ContentBlock(
            type=BlockType.PARAGRAPH,
            position=Position(page=1, index=0),
            text="dead",
            runs=[TextRun("dead", struck=True)],
        )
        assert b.live_text() == ""

    def test_legacy_block_no_runs_falls_back_to_text(self):
        b = ContentBlock(
            type=BlockType.PARAGRAPH,
            position=Position(page=1, index=0),
            text="legacy",
            font_info=FontInfo(size=11.0, strikethrough=False),
        )
        assert b.live_text() == "legacy"

    def test_legacy_block_struck_returns_empty(self):
        b = ContentBlock(
            type=BlockType.PARAGRAPH,
            position=Position(page=1, index=0),
            text="legacy",
            font_info=FontInfo(size=11.0, strikethrough=True),
        )
        assert b.live_text() == ""


class TestRowStruck:
    def _table(self, header_runs, row_runs):
        return ContentBlock(
            type=BlockType.TABLE,
            position=Position(page=1, index=0),
            headers=["A", "B"],
            rows=[["", ""] for _ in row_runs],
            header_runs=header_runs,
            row_runs=row_runs,
        )

    def test_row_all_struck_when_every_textful_cell_fully_struck(self):
        b = self._table(
            header_runs=[[TextRun("A")], [TextRun("B")]],
            row_runs=[
                [[TextRun("22")], [TextRun("exponential")]],
                [[TextRun("23", struck=True)], [TextRun("stop", struck=True)]],
            ],
        )
        assert not b.row_all_struck(0)
        assert b.row_all_struck(1)

    def test_partial_cell_strike_does_not_mark_row_struck(self):
        # One cell has mixed runs (typo correction). Row should NOT be
        # considered struck — strict criterion per user direction.
        b = self._table(
            header_runs=[],
            row_runs=[
                [
                    [TextRun("Code "), TextRun("22", struck=True), TextRun(" 24")],
                    [TextRun("exponential")],
                ],
            ],
        )
        assert not b.row_all_struck(0)

    def test_header_all_struck(self):
        b = self._table(
            header_runs=[[TextRun("A", struck=True)], [TextRun("B", struck=True)]],
            row_runs=[],
        )
        assert b.header_all_struck()

    def test_cell_live_text_drops_struck_runs(self):
        b = self._table(
            header_runs=[],
            row_runs=[
                [[TextRun("Code "), TextRun("22", struck=True), TextRun(" 24")]],
            ],
        )
        assert b.cell_live_text(0, 0) == "Code  24"


class TestJSONRoundtrip:
    def test_runs_survive_save_load(self, tmp_path: Path):
        b = ContentBlock(
            type=BlockType.PARAGRAPH,
            position=Position(page=1, index=0),
            text="hello world",
            runs=[
                TextRun("hello ", struck=False),
                TextRun("world", struck=True),
            ],
        )
        ir = DocumentIR(
            source_file="t.docx", source_format="docx", content_blocks=[b]
        )
        path = tmp_path / "ir.json"
        ir.save_json(path)
        ir2 = DocumentIR.load_json(path)
        assert len(ir2.content_blocks[0].runs) == 2
        assert ir2.content_blocks[0].live_text() == "hello "

    def test_table_runs_survive_save_load(self, tmp_path: Path):
        b = ContentBlock(
            type=BlockType.TABLE,
            position=Position(page=1, index=0),
            headers=["A", "B"],
            rows=[["1", "2"]],
            header_runs=[[TextRun("A")], [TextRun("B")]],
            row_runs=[[[TextRun("1", struck=True)], [TextRun("2", struck=True)]]],
        )
        ir = DocumentIR(
            source_file="t.docx", source_format="docx", content_blocks=[b]
        )
        path = tmp_path / "ir.json"
        ir.save_json(path)
        ir2 = DocumentIR.load_json(path)
        assert ir2.content_blocks[0].row_all_struck(0)


# ---------------------------------------------------------------------------
# DOCX extractor — partial paragraph + struck table rows
# ---------------------------------------------------------------------------

def _build_partial_strike_docx(path: Path) -> None:
    doc = DocxDocument()
    # Partial-strike paragraph
    p = doc.add_paragraph()
    p.add_run("The UE shall ")
    s = p.add_run("not "); s.font.strike = True
    p.add_run("retry within 30s.")
    # Fully-struck heading
    h = doc.add_heading("", level=1)
    hr = h.add_run("Withdrawn Section"); hr.font.strike = True
    # Body paragraph (would cascade with the struck heading at parse time,
    # but the extractor itself just keeps it).
    doc.add_paragraph("Body under withdrawn.")
    # Table with a fully-struck second body row
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Code"
    t.cell(0, 1).text = "Action"
    t.cell(1, 0).text = "22"
    t.cell(1, 1).text = "exponential"
    # Make row 2 fully struck via runs
    for col in (0, 1):
        cell = t.cell(2, col)
        cell.text = ""  # clear default empty paragraph
        run = cell.paragraphs[0].add_run("23" if col == 0 else "stop")
        run.font.strike = True
    doc.save(str(path))


class TestDocxExtractorRuns:
    @pytest.fixture
    def doc_path(self, tmp_path: Path) -> Path:
        path = tmp_path / "fixture.docx"
        _build_partial_strike_docx(path)
        return path

    def test_partial_paragraph_runs_preserved(self, doc_path: Path):
        ir = DOCXExtractor().extract(doc_path)
        para = next(
            b for b in ir.content_blocks
            if b.type == BlockType.PARAGRAPH and "retry" in b.text
        )
        assert len(para.runs) == 3
        assert [r.struck for r in para.runs] == [False, True, False]
        assert para.live_text() == "The UE shall retry within 30s."
        # Partial strike → block-level flag is False
        assert para.font_info.strikethrough is False

    def test_full_strike_heading_marked_block_level(self, doc_path: Path):
        ir = DOCXExtractor().extract(doc_path)
        heading = next(
            b for b in ir.content_blocks
            if b.type == BlockType.HEADING and b.text == "Withdrawn Section"
        )
        assert heading.font_info.strikethrough is True
        assert heading.live_text() == ""

    def test_struck_table_row_kept_with_row_runs_marked(self, doc_path: Path):
        ir = DOCXExtractor().extract(doc_path)
        table = next(b for b in ir.content_blocks if b.type == BlockType.TABLE)
        # Both body rows preserved (no drop at extract time)
        assert len(table.rows) == 2
        # Row 0 is live; row 1 is fully struck per row_runs
        assert not table.row_all_struck(0)
        assert table.row_all_struck(1)
        # The whole table is NOT struck (header + row 0 are live)
        assert table.font_info is None or table.font_info.strikethrough is False


# ---------------------------------------------------------------------------
# Parser — partial-strike normalization
# ---------------------------------------------------------------------------

class TestParserPartialStrike:
    """Parser drops struck spans + struck rows when ignore_strikeout=True."""

    def _parse(self, ir: DocumentIR):
        from core.src.profiler.profile_schema import (
            DocumentProfile,
            HeadingDetection,
            RequirementIdPattern,
        )
        from core.src.parser.structural_parser import GenericStructuralParser

        # Minimal profile that won't classify our test paragraphs as
        # headings — just need ignore_strikeout default (True).
        profile = DocumentProfile(
            profile_name="test",
            heading_detection=HeadingDetection(),
            requirement_id=RequirementIdPattern(pattern="VZ_REQ_[A-Z0-9_]+"),
        )
        return GenericStructuralParser(profile).parse(ir)

    def test_partial_paragraph_text_normalized_to_live(self):
        # A standalone paragraph with partial strike. The parser walks
        # the IR; for non-fully-struck blocks with runs, it normalizes
        # block.text to live_text.
        ir = DocumentIR(
            source_file="t.docx",
            source_format="docx",
            content_blocks=[
                ContentBlock(
                    type=BlockType.PARAGRAPH,
                    position=Position(page=1, index=0),
                    text="The UE shall not retry within 30s.",
                    runs=[
                        TextRun("The UE shall ", struck=False),
                        TextRun("not ", struck=True),
                        TextRun("retry within 30s.", struck=False),
                    ],
                    font_info=FontInfo(size=11.0, strikethrough=False),
                ),
            ],
        )
        # We don't assert on the parsed-tree shape (depends on a lot of
        # profile config); we assert the block was mutated to live_text.
        self._parse(ir)
        assert ir.content_blocks[0].text == "The UE shall retry within 30s."

    def test_fully_struck_row_dropped_from_table_block(self):
        ir = DocumentIR(
            source_file="t.docx",
            source_format="docx",
            content_blocks=[
                ContentBlock(
                    type=BlockType.TABLE,
                    position=Position(page=1, index=0),
                    headers=["Code", "Action"],
                    rows=[["22", "exponential"], ["23", "stop"]],
                    header_runs=[[TextRun("Code")], [TextRun("Action")]],
                    row_runs=[
                        [[TextRun("22")], [TextRun("exponential")]],
                        [[TextRun("23", struck=True)],
                         [TextRun("stop", struck=True)]],
                    ],
                ),
            ],
        )
        self._parse(ir)
        # Parser mutates rows in place: struck row dropped
        assert ir.content_blocks[0].rows == [["22", "exponential"]]

    def test_docx_heading_cascade_arms_on_struck_block_level(self):
        """D-061: _heading_depth honors BlockType.HEADING with block.level
        (DOCX-style headings) so the FR-33 cascade arms when a struck
        DOCX heading appears. Pre-D-061 it only fired for PARAGRAPH-typed
        headings (PDF convention) — silently broken for DOCX.
        """
        ir = DocumentIR(
            source_file="t.docx",
            source_format="docx",
            content_blocks=[
                ContentBlock(
                    type=BlockType.HEADING,
                    position=Position(page=1, index=0),
                    text="Withdrawn Section",
                    level=1,
                    runs=[TextRun("Withdrawn Section", struck=True)],
                    font_info=FontInfo(size=18.0, strikethrough=True),
                ),
                # Body paragraph that should cascade-drop
                ContentBlock(
                    type=BlockType.PARAGRAPH,
                    position=Position(page=1, index=1),
                    text="Body under withdrawn section.",
                    runs=[TextRun("Body under withdrawn section.")],
                    font_info=FontInfo(size=11.0, strikethrough=False),
                ),
                # Sibling depth-1 heading should NOT be cascaded
                ContentBlock(
                    type=BlockType.HEADING,
                    position=Position(page=1, index=2),
                    text="Live Section",
                    level=1,
                    runs=[TextRun("Live Section")],
                    font_info=FontInfo(size=18.0, strikethrough=False),
                ),
                ContentBlock(
                    type=BlockType.PARAGRAPH,
                    position=Position(page=1, index=3),
                    text="Live body content.",
                    runs=[TextRun("Live body content.")],
                    font_info=FontInfo(size=11.0, strikethrough=False),
                ),
            ],
        )
        tree = self._parse(ir)
        # Cascade fired: parse_stats.cascade_blocks_dropped > 0
        assert tree.parse_stats.cascade_blocks_dropped >= 1

    def test_struck_req_id_in_partial_span_removed_from_block_text(self):
        # Observable effect of the partial-strike normalization: the
        # struck req_id token is removed from block.text. The internal
        # struck_req_ids set (used as a table-anchored skip set) is not
        # part of RequirementTree's public surface, but the visible
        # downstream is that the req_id no longer appears in any chunk
        # the parser feeds to the rest of the pipeline.
        ir = DocumentIR(
            source_file="t.docx",
            source_format="docx",
            content_blocks=[
                ContentBlock(
                    type=BlockType.PARAGRAPH,
                    position=Position(page=1, index=0),
                    text="VZ_REQ_LTEAT_99 still applies",
                    runs=[
                        TextRun("VZ_REQ_LTEAT_99 ", struck=True),
                        TextRun("still applies", struck=False),
                    ],
                    font_info=FontInfo(size=11.0, strikethrough=False),
                ),
            ],
        )
        self._parse(ir)
        assert "VZ_REQ_LTEAT_99" not in ir.content_blocks[0].text
        assert ir.content_blocks[0].text == "still applies"
