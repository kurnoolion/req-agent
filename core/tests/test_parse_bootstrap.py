"""Tests for the Bootstrap annotation harness.

Three layers:
  1. ``bootstrap_schema.validate_annotation_file`` — schema enforcement.
  2. ``docx_html_render.render_docx_html`` — block-index alignment with
     :class:`DOCXExtractor`.
  3. Routes — list docs, view, load/save annotations via TestClient.
"""

from __future__ import annotations

import importlib
import json
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from core.src.extraction.docx_extractor import DOCXExtractor
from core.src.web.bootstrap_schema import (
    AnnotationValidationError,
    SCHEMA_VERSION,
    validate_annotation_file,
)
from core.src.web.docx_html_render import render_docx_html


# ---------------------------------------------------------------------------
# Schema validation
# ---------------------------------------------------------------------------

class TestSchemaValidation:
    def _payload(self, **anns) -> dict:
        return {
            "version": SCHEMA_VERSION,
            "doc_path": "<env_dir>/input/MNO/REL/plan.docx",
            "annotations": list(anns.get("annotations", [])),
        }

    def test_minimal_section_heading_passes(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "section_heading", "region": {"block_indices": [42]}},
        ]))
        assert out["annotations"][0]["region"] == {"block_indices": [42]}

    def test_block_index_with_row_range(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "strikethrough",
             "region": {"block_index": 10, "row_range": [2, 5]},
             "subkind": "table_row"},
        ]))
        assert out["annotations"][0]["region"] == {"block_index": 10, "row_range": [2, 5]}
        assert out["annotations"][0]["subkind"] == "table_row"

    def test_unknown_kind_rejected(self):
        with pytest.raises(AnnotationValidationError) as exc:
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "not_a_kind", "region": {"block_indices": [1]}},
            ]))
        assert any("kind" in e for e in exc.value.errors)

    def test_old_references_kind_rejected(self):
        # The old single `references` kind is gone; split into 3 kinds.
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "references", "region": {"block_indices": [3]}},
            ]))

    def test_reference_intra_doc_minimal(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_intra_doc",
             "region": {"block_indices": [3]}},
        ]))
        assert out["annotations"][0]["kind"] == "reference_intra_doc"

    def test_reference_intra_doc_with_target(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_intra_doc",
             "region": {"block_indices": [3]},
             "inline": True,
             "target": {"section_number": "3.5.2.1"}},
        ]))
        ann = out["annotations"][0]
        assert ann["target"] == {"section_number": "3.5.2.1"}
        assert ann["inline"] is True

    def test_reference_cross_doc_with_target(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_cross_doc",
             "region": {"block_indices": [3]},
             "target": {"plan_id": "<PLAN1>", "req_id": "<MNO>_REQ_<PLAN1>_45"}},
        ]))
        ann = out["annotations"][0]
        assert ann["target"]["plan_id"] == "<PLAN1>"
        assert ann["target"]["req_id"] == "<MNO>_REQ_<PLAN1>_45"

    def test_reference_spec_requires_style(self):
        with pytest.raises(AnnotationValidationError) as exc:
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "reference_spec",
                 "region": {"block_indices": [3]}},
            ]))
        assert any("style" in e for e in exc.value.errors)

    def test_reference_spec_direct_with_target(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_spec", "style": "direct",
             "region": {"block_indices": [3]},
             "target": {"spec": "3GPP TS 24.301", "section": "5.5.1.2.6"}},
        ]))
        ann = out["annotations"][0]
        assert ann["style"] == "direct"
        assert ann["target"] == {"spec": "3GPP TS 24.301", "section": "5.5.1.2.6"}

    def test_reference_spec_indirect_with_target(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_spec", "style": "indirect",
             "region": {"block_indices": [3]},
             "target": {"ref_number": 5}},
        ]))
        ann = out["annotations"][0]
        assert ann["style"] == "indirect"
        assert ann["target"] == {"ref_number": 5}

    def test_reference_spec_invalid_style_rejected(self):
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "reference_spec", "style": "fancy",
                 "region": {"block_indices": [3]}},
            ]))

    def test_reference_list_with_layout(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_list",
             "region": {"block_indices": [201, 202, 203]},
             "numbering_style": "bracketed", "layout": "paragraph_list"},
        ]))
        ann = out["annotations"][0]
        assert ann["numbering_style"] == "bracketed"
        assert ann["layout"] == "paragraph_list"

    def test_reference_list_entry_with_number_and_target(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_list_entry",
             "region": {"block_indices": [205]},
             "number": 5, "title_hint_chars": 67,
             "target": {"spec": "3GPP TS 24.301"}},
        ]))
        ann = out["annotations"][0]
        assert ann["number"] == 5
        assert ann["title_hint_chars"] == 67
        assert ann["target"] == {"spec": "3GPP TS 24.301"}

    def test_target_unknown_keys_stripped(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "reference_intra_doc",
             "region": {"block_indices": [3]},
             "target": {"section_number": "1.2", "future_field": "drop"}},
        ]))
        assert "future_field" not in out["annotations"][0]["target"]
        assert out["annotations"][0]["target"]["section_number"] == "1.2"

    def test_target_ref_number_must_be_int(self):
        with pytest.raises(AnnotationValidationError) as exc:
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "reference_spec", "style": "indirect",
                 "region": {"block_indices": [3]},
                 "target": {"ref_number": "five"}},
            ]))
        assert any("ref_number" in e for e in exc.value.errors)

    def test_target_must_be_object_not_list(self):
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "reference_intra_doc",
                 "region": {"block_indices": [3]},
                 "target": ["not", "a", "dict"]},
            ]))

    def test_block_indices_must_be_non_negative_ints(self):
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "toc", "region": {"block_indices": [-1]}},
            ]))
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "toc", "region": {"block_indices": []}},
            ]))

    def test_row_range_must_be_ordered(self):
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "strikethrough",
                 "region": {"block_index": 1, "row_range": [5, 3]}},
            ]))

    def test_notes_length_capped(self):
        long_note = "x" * 31
        with pytest.raises(AnnotationValidationError) as exc:
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "toc", "region": {"block_indices": [1]},
                 "notes": long_note},
            ]))
        assert any("notes" in e for e in exc.value.errors)

    def test_duplicate_ids_rejected(self):
        with pytest.raises(AnnotationValidationError) as exc:
            validate_annotation_file(self._payload(annotations=[
                {"id": "ann_001", "kind": "toc", "region": {"block_indices": [1]}},
                {"id": "ann_001", "kind": "toc", "region": {"block_indices": [2]}},
            ]))
        assert any("duplicate" in e for e in exc.value.errors)

    def test_section_heading_depth_bounds(self):
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "section_heading",
                 "region": {"block_indices": [1]}, "depth": 0},
            ]))
        with pytest.raises(AnnotationValidationError):
            validate_annotation_file(self._payload(annotations=[
                {"id": "x", "kind": "section_heading",
                 "region": {"block_indices": [1]}, "depth": 10},
            ]))

    def test_unknown_optional_fields_stripped(self):
        out = validate_annotation_file(self._payload(annotations=[
            {"id": "ann_001", "kind": "toc",
             "region": {"block_indices": [1]},
             "future_field": "drop me"},
        ]))
        assert "future_field" not in out["annotations"][0]


# ---------------------------------------------------------------------------
# DOCX renderer alignment with DOCXExtractor
# ---------------------------------------------------------------------------

def _build_fixture_docx(path: Path) -> None:
    """Minimal DOCX exercising heading, paragraph, empty paragraph, table.

    The empty paragraph is critical — DOCXExtractor skips it, and so must
    the renderer, otherwise data-block-idx values drift.
    """
    doc = DocxDocument()
    doc.add_heading("Chapter 1", level=1)         # block 0 (HEADING)
    doc.add_paragraph("Body text under chapter.")  # block 1 (PARAGRAPH)
    doc.add_paragraph("")                          # SKIPPED — empty
    doc.add_heading("Section 1.1", level=2)        # block 2 (HEADING)
    doc.add_paragraph("Another paragraph.")        # block 3 (PARAGRAPH)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Code"
    table.cell(0, 1).text = "Action"
    table.cell(1, 0).text = "22"
    table.cell(1, 1).text = "exponential"
    table.cell(2, 0).text = "23"
    table.cell(2, 1).text = "stop"                 # block 4 (TABLE)
    doc.save(str(path))


class TestDocxRenderAlignment:
    def test_block_indices_match_extractor(self, tmp_path):
        docx_path = tmp_path / "fixture.docx"
        _build_fixture_docx(docx_path)

        ir = DOCXExtractor().extract(docx_path)
        ir_indices = [b.position.index for b in ir.content_blocks]
        assert ir_indices == [0, 1, 2, 3, 4]

        html = render_docx_html(docx_path)
        # Every IR block index appears as a data-block-idx attribute.
        for idx in ir_indices:
            assert f'data-block-idx="{idx}"' in html, (
                f"missing data-block-idx={idx} in html:\n{html}"
            )
        # No extra index emitted (e.g. for the empty paragraph).
        assert f'data-block-idx="{len(ir_indices)}"' not in html

    def test_table_row_indices_emitted(self, tmp_path):
        docx_path = tmp_path / "fixture.docx"
        _build_fixture_docx(docx_path)
        html = render_docx_html(docx_path)
        assert 'data-row-idx="0"' in html  # body row 0 (after header)
        assert 'data-row-idx="1"' in html  # body row 1
        # No row-idx on the header row.
        assert 'data-row-idx="header"' not in html

    def test_heading_emits_h_tag(self, tmp_path):
        docx_path = tmp_path / "fixture.docx"
        _build_fixture_docx(docx_path)
        html = render_docx_html(docx_path)
        assert "<h1" in html
        assert "<h2" in html
        assert "Chapter 1" in html


# ---------------------------------------------------------------------------
# Route tests
# ---------------------------------------------------------------------------

@pytest.fixture
def web_client(tmp_path, monkeypatch):
    """TestClient bound to a tmp env_dir."""
    from fastapi.testclient import TestClient
    monkeypatch.setenv("ENV_DIR", str(tmp_path))
    monkeypatch.delenv("NORA_CONFIG_DB", raising=False)
    import core.src.web.app as app_mod
    importlib.reload(app_mod)
    return TestClient(app_mod.app), tmp_path


class TestBootstrapRoutes:
    def test_list_docs_empty(self, web_client):
        client, _ = web_client
        resp = client.get("/parse-review/bootstrap/docs")
        assert resp.status_code == 200
        body = resp.json()
        assert body["docs"] == []
        assert "annotations" in body["annotations_dir"]

    def test_list_docs_with_input(self, web_client):
        client, env_dir = web_client
        plan_dir = env_dir / "input" / "VZW" / "OA-test"
        plan_dir.mkdir(parents=True)
        docx_path = plan_dir / "LTEAT.docx"
        _build_fixture_docx(docx_path)
        # No IR yet → ir_exists False
        resp = client.get("/parse-review/bootstrap/docs")
        body = resp.json()
        assert len(body["docs"]) == 1
        assert body["docs"][0]["doc_id"] == "LTEAT"
        assert body["docs"][0]["ir_exists"] is False

        # Run extract → IR exists → ir_exists True
        ir = DOCXExtractor().extract(docx_path)
        ir_dir = env_dir / "out" / "extract"
        ir_dir.mkdir(parents=True)
        ir.save_json(ir_dir / "LTEAT_ir.json")
        resp = client.get("/parse-review/bootstrap/docs")
        body = resp.json()
        assert body["docs"][0]["ir_exists"] is True

    def test_view_renders_blocks(self, web_client):
        client, env_dir = web_client
        plan_dir = env_dir / "input" / "VZW" / "OA-test"
        plan_dir.mkdir(parents=True)
        docx_path = plan_dir / "LTEAT.docx"
        _build_fixture_docx(docx_path)
        ir = DOCXExtractor().extract(docx_path)
        (env_dir / "out" / "extract").mkdir(parents=True)
        ir.save_json(env_dir / "out" / "extract" / "LTEAT_ir.json")

        resp = client.get("/parse-review/bootstrap/LTEAT/view")
        assert resp.status_code == 200
        # IR-block markers
        assert 'data-idx="0"' in resp.text
        assert 'data-idx="4"' in resp.text  # the table block
        # DOCX preview markers (renderer output)
        assert 'data-block-idx="0"' in resp.text
        # Bootstrap save button
        assert "bs-save-btn" in resp.text

    def test_load_annotations_when_absent(self, web_client):
        client, env_dir = web_client
        plan_dir = env_dir / "input" / "VZW" / "OA-test"
        plan_dir.mkdir(parents=True)
        _build_fixture_docx(plan_dir / "LTEAT.docx")
        resp = client.get("/parse-review/bootstrap/LTEAT/annotations")
        assert resp.status_code == 200
        body = resp.json()
        assert body["annotations"] == []
        assert body["doc_path"].endswith("LTEAT.docx")

    def test_save_then_load_roundtrip(self, web_client):
        client, env_dir = web_client
        plan_dir = env_dir / "input" / "VZW" / "OA-test"
        plan_dir.mkdir(parents=True)
        _build_fixture_docx(plan_dir / "LTEAT.docx")
        payload = {
            "version": SCHEMA_VERSION,
            "doc_path": "<env_dir>/input/VZW/OA-test/LTEAT.docx",
            "annotations": [
                {"id": "ann_001", "kind": "section_heading",
                 "region": {"block_indices": [0]}, "depth": 1, "section_number": "1"},
                {"id": "ann_002", "kind": "strikethrough",
                 "region": {"block_index": 4, "row_range": [0, 0]},
                 "subkind": "table_row"},
                {"id": "ann_003", "kind": "reference_spec", "style": "indirect",
                 "region": {"block_indices": [3]},
                 "target": {"ref_number": 5}},
            ],
        }
        resp = client.post("/parse-review/bootstrap/LTEAT/annotations", json=payload)
        assert resp.status_code == 200, resp.text
        assert resp.json()["ok"] is True

        path = env_dir / "annotations" / "LTEAT_annotations.json"
        assert path.exists()
        on_disk = json.loads(path.read_text(encoding="utf-8"))
        assert len(on_disk["annotations"]) == 3

        resp = client.get("/parse-review/bootstrap/LTEAT/annotations")
        body = resp.json()
        assert len(body["annotations"]) == 3
        assert body["annotations"][1]["region"]["row_range"] == [0, 0]
        assert body["annotations"][2]["target"] == {"ref_number": 5}

    def test_save_validation_error_returns_400(self, web_client):
        client, env_dir = web_client
        plan_dir = env_dir / "input" / "VZW" / "OA-test"
        plan_dir.mkdir(parents=True)
        _build_fixture_docx(plan_dir / "LTEAT.docx")
        bad_payload = {
            "version": SCHEMA_VERSION,
            "doc_path": "<env_dir>/.../LTEAT.docx",
            "annotations": [
                {"id": "ann_001", "kind": "reference_spec",  # missing style
                 "region": {"block_indices": [1]}},
            ],
        }
        resp = client.post("/parse-review/bootstrap/LTEAT/annotations", json=bad_payload)
        assert resp.status_code == 400
        body = resp.json()
        assert body["ok"] is False
        assert any("style" in e for e in body["errors"])
