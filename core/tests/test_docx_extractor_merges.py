"""Tests for the DOCX extractor's merged-cell metadata.

Builds in-memory .docx files via python-docx with deliberately merged
table regions, runs the extractor, and asserts that the resulting
ContentBlock has the expected `merged_cells` shape AND that continuation
positions in the rectangular `headers`/`rows` matrices are blanked out
(no duplicate-text artefact from python-docx's `row.cells` API).
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document  # noqa: E402

from core.src.extraction.docx_extractor import DOCXExtractor  # noqa: E402


def _extract(path: Path):
    return DOCXExtractor().extract(path)


def _tables(ir):
    return [b for b in ir.content_blocks if b.type.value == "table"]


# ---------------------------------------------------------------------------
# Horizontal merge (gridSpan)
# ---------------------------------------------------------------------------

def test_horizontal_merge_anchor_in_header_row(tmp_path: Path):
    """Header row's first two cells are merged horizontally. The merged
    cell carries the label 'Revision History'. Body has 1 row, 3 cells."""
    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Rev."
    tbl.rows[0].cells[1].text = "Author"
    tbl.rows[0].cells[2].text = "Date"
    # Merge header cells 0..1 — gridSpan=2
    merged_anchor = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    merged_anchor.text = "Revision History"
    # Body row
    tbl.rows[1].cells[0].text = "1.0"
    tbl.rows[1].cells[1].text = "Alice"
    tbl.rows[1].cells[2].text = "2026-01-01"

    p = tmp_path / "h.docx"
    doc.save(p)
    ir = _extract(p)
    tables = _tables(ir)
    assert len(tables) == 1
    t = tables[0]

    # Merge metadata: one entry, anchor at (0,0), colspan=2.
    assert len(t.merged_cells) == 1
    mc = t.merged_cells[0]
    assert (mc.row, mc.col, mc.rowspan, mc.colspan) == (0, 0, 1, 2)
    assert mc.text == "Revision History"

    # Header matrix: anchor text at col 0, continuation BLANK at col 1.
    assert t.headers == ["Revision History", "", "Date"]
    # Body row unaffected by header merge.
    assert t.rows == [["1.0", "Alice", "2026-01-01"]]


def test_horizontal_merge_anchor_in_last_row(tmp_path: Path):
    """Footer-style merged row at the bottom of the table — common
    'Revision History' / 'Document Control' label in proprietary forms."""
    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    tbl.rows[1].cells[0].text = "Version"
    tbl.rows[1].cells[1].text = "2.0"
    # Merge bottom row across all columns.
    anchor = tbl.rows[2].cells[0].merge(tbl.rows[2].cells[1])
    anchor.text = "Revision History"

    p = tmp_path / "h2.docx"
    doc.save(p)
    ir = _extract(p)
    t = _tables(ir)[0]
    assert len(t.merged_cells) == 1
    mc = t.merged_cells[0]
    # Header row counts as row 0; bottom row is row 2.
    assert (mc.row, mc.col, mc.rowspan, mc.colspan) == (2, 0, 1, 2)
    assert mc.text == "Revision History"


# ---------------------------------------------------------------------------
# Vertical merge (vMerge)
# ---------------------------------------------------------------------------

def test_vertical_merge_anchor_at_top(tmp_path: Path):
    """Col 0 rows 0..2 merged vertically — anchor in row 0."""
    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Section"
    tbl.rows[0].cells[1].text = "Desc"
    tbl.rows[1].cells[0].text = "X"   # will be overridden by merge
    tbl.rows[1].cells[1].text = "row1"
    tbl.rows[2].cells[0].text = "Y"
    tbl.rows[2].cells[1].text = "row2"
    # Merge col 0 across all 3 rows.
    anchor = tbl.rows[0].cells[0].merge(tbl.rows[1].cells[0]).merge(
        tbl.rows[2].cells[0]
    )
    anchor.text = "Same Section"

    p = tmp_path / "v.docx"
    doc.save(p)
    ir = _extract(p)
    t = _tables(ir)[0]
    assert len(t.merged_cells) == 1
    mc = t.merged_cells[0]
    assert (mc.row, mc.col, mc.rowspan, mc.colspan) == (0, 0, 3, 1)
    assert mc.text == "Same Section"

    # Header row col 0 carries the anchor text. Body rows (which are
    # rows[1] and rows[2] in the IR) carry "" at col 0.
    assert t.headers == ["Same Section", "Desc"]
    assert t.rows == [["", "row1"], ["", "row2"]]


# ---------------------------------------------------------------------------
# Non-merged tables — no regression
# ---------------------------------------------------------------------------

def test_non_merged_table_emits_empty_merged_cells(tmp_path: Path):
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[1].cells[0].text = "C"
    tbl.rows[1].cells[1].text = "D"

    p = tmp_path / "plain.docx"
    doc.save(p)
    ir = _extract(p)
    t = _tables(ir)[0]
    assert t.merged_cells == []
    assert t.headers == ["A", "B"]
    assert t.rows == [["C", "D"]]


# ---------------------------------------------------------------------------
# Multiple independent merges in one table
# ---------------------------------------------------------------------------

def test_multiple_merges_recorded(tmp_path: Path):
    doc = Document()
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    # Header row: merge cols 0+1
    tbl.rows[0].cells[2].text = "C"
    a1 = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    a1.text = "AB"
    # Body row 1: merge cols 1+2
    tbl.rows[1].cells[0].text = "x"
    a2 = tbl.rows[1].cells[1].merge(tbl.rows[1].cells[2])
    a2.text = "yz"
    # Body row 2: no merge
    tbl.rows[2].cells[0].text = "p"
    tbl.rows[2].cells[1].text = "q"
    tbl.rows[2].cells[2].text = "r"

    p = tmp_path / "multi.docx"
    doc.save(p)
    ir = _extract(p)
    t = _tables(ir)[0]
    assert len(t.merged_cells) == 2
    # Sorted by (row, col) ⇒ header-row merge first.
    assert (t.merged_cells[0].row, t.merged_cells[0].col,
            t.merged_cells[0].colspan) == (0, 0, 2)
    assert (t.merged_cells[1].row, t.merged_cells[1].col,
            t.merged_cells[1].colspan) == (1, 1, 2)


# ---------------------------------------------------------------------------
# JSON round-trip
# ---------------------------------------------------------------------------

def test_merged_cells_survive_json_roundtrip(tmp_path: Path):
    from core.src.models.document import DocumentIR

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "x"
    tbl.rows[0].cells[1].text = "y"
    anchor = tbl.rows[1].cells[0].merge(tbl.rows[1].cells[1])
    anchor.text = "Footer Label"

    p = tmp_path / "rt.docx"
    doc.save(p)
    ir = _extract(p)

    out = tmp_path / "rt_ir.json"
    ir.save_json(out)
    reloaded = DocumentIR.load_json(out)
    t = _tables(reloaded)[0]
    assert len(t.merged_cells) == 1
    assert t.merged_cells[0].text == "Footer Label"
    assert t.merged_cells[0].colspan == 2


# ---------------------------------------------------------------------------
# Degenerate-table filter — drops only when every cell is empty.
# ---------------------------------------------------------------------------

def test_single_cell_content_table_preserved(tmp_path: Path):
    """Word documents commonly use 1×1 tables as paragraph containers
    (e.g. a section's entire body inside a one-cell table for layout
    purposes). The prior filter shape (`non_empty_headers <= 1 and
    total_cells == 0`) accidentally dropped these — the single content
    cell becomes the only "header", body is empty, both conditions
    trigger. Real corpus regression: a doc's next-section content was
    missing from the IR after a glossary section because it was wrapped
    in a 1×1 table."""
    doc = Document()
    doc.add_paragraph("Some intro paragraph.")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Real content trapped in a 1x1 wrapper table."

    p = tmp_path / "single_cell.docx"
    doc.save(p)
    ir = _extract(p)
    tables = _tables(ir)
    assert len(tables) == 1
    assert tables[0].headers == [
        "Real content trapped in a 1x1 wrapper table.",
    ]
    assert tables[0].rows == []


def test_empty_table_still_dropped(tmp_path: Path):
    """The filter's actual intent — a fully-empty table contributes
    nothing and is still dropped. Guards against the loosening
    over-shooting."""
    doc = Document()
    doc.add_paragraph("Before.")
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    # leave every cell empty

    p = tmp_path / "empty.docx"
    doc.save(p)
    ir = _extract(p)
    assert _tables(ir) == []


def test_nested_glossary_table_inside_wrapper_emitted_separately(
    tmp_path: Path,
):
    """Real corpus pattern: an outer 1×1 wrapper table whose single
    cell contains both a nested 2-column glossary table AND a trailing
    paragraph. Before the nested walk landed, the nested table was
    invisible — `cell.text` returns only paragraph text, so the
    acronym/definition rows never reached the IR. The parser then
    matched the glossary section by title but found no entries.

    After the fix: two TABLE blocks land in the IR — outer wrapper
    carries the trailing paragraph text (its only readable
    `cell.text` content), nested carries the actual headers + rows."""
    doc = Document()
    doc.add_paragraph("1.1 Glossary")
    outer = doc.add_table(rows=1, cols=1)
    outer.style = "Table Grid"
    cell = outer.rows[0].cells[0]
    # Nested glossary table at the top of the cell.
    nested = cell.add_table(rows=3, cols=2)
    nested.style = "Table Grid"
    nested.rows[0].cells[0].text = "Acronym/Term"
    nested.rows[0].cells[1].text = "Definition"
    nested.rows[1].cells[0].text = "ETWS"
    nested.rows[1].cells[1].text = "Earthquake and Tsunami Warning System"
    nested.rows[2].cells[0].text = "APN"
    nested.rows[2].cells[1].text = "Access Point Name"
    # Trailing paragraph in the same cell, below the nested table.
    cell.add_paragraph("Trailing description under the glossary table.")

    p = tmp_path / "nested.docx"
    doc.save(p)
    ir = _extract(p)
    tables = _tables(ir)
    assert len(tables) == 2

    # Outer wrapper sees only the trailing paragraph in `cell.text`.
    outer_block = tables[0]
    assert outer_block.headers == [
        "Trailing description under the glossary table.",
    ]

    # Nested glossary table carries the real content.
    nested_block = tables[1]
    assert nested_block.headers == ["Acronym/Term", "Definition"]
    assert nested_block.rows == [
        ["ETWS", "Earthquake and Tsunami Warning System"],
        ["APN", "Access Point Name"],
    ]


def test_nested_table_emitted_even_when_outer_is_empty_wrapper(
    tmp_path: Path,
):
    """When the outer 1×1 wrapper is a pure layout container — no
    surrounding text in the cell, only a nested table — the wrapper
    block itself is correctly dropped by the empty-table filter, but
    the nested table must still survive."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    outer.style = "Table Grid"
    cell = outer.rows[0].cells[0]
    # The empty paragraph that python-docx auto-inserts in a new cell
    # has no text — outer.cell.text is empty.
    nested = cell.add_table(rows=2, cols=2)
    nested.style = "Table Grid"
    nested.rows[0].cells[0].text = "Term"
    nested.rows[0].cells[1].text = "Definition"
    nested.rows[1].cells[0].text = "RAT"
    nested.rows[1].cells[1].text = "Radio Access Technology"

    p = tmp_path / "wrapper_only.docx"
    doc.save(p)
    ir = _extract(p)
    tables = _tables(ir)
    # Outer wrapper dropped (empty); nested survives.
    assert len(tables) == 1
    assert tables[0].headers == ["Term", "Definition"]
    assert tables[0].rows == [["RAT", "Radio Access Technology"]]


def test_sparse_table_with_one_content_cell_preserved(tmp_path: Path):
    """1×N table where only one cell has content — also was dropped by
    the prior filter. Now survives so the parser sees the content."""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[1].text = "Lone content"

    p = tmp_path / "sparse.docx"
    doc.save(p)
    ir = _extract(p)
    tables = _tables(ir)
    assert len(tables) == 1
    assert tables[0].headers == ["", "Lone content", ""]
