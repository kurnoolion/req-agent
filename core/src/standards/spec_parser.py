"""3GPP specification document parser (DOC/DOCX to section tree).

Parses 3GPP specification documents into a structured section tree
using paragraph styles (Heading 1-9) and section numbering. Rule-based,
no LLM required.

Handles:
- Numbered sections (e.g., 5.5.1.2.5)
- Lettered sub-sections (e.g., 4.2A)
- Annexes (e.g., Annex A)
- Foreword and unnumbered front matter
- Tab-separated section numbers and titles
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

import docx

from src.standards.schema import SpecDocument, SpecSection

logger = logging.getLogger(__name__)

# Section number pattern: "5.5.1.2.5" or "4.2A" (with optional letter suffix)
_SECTION_NUM = re.compile(r"^([\d]+(?:\.[\d]+)*[A-Za-z]?)\t(.*)$")

# Annex pattern: "Annex A (informative): Description"
_ANNEX_PAT = re.compile(
    r"^(Annex\s+[A-Z][A-Za-z0-9]*)\s*[\(:]?\s*(.*?)[\):]?\s*$",
    re.IGNORECASE,
)

# Spec title in front matter: "3GPP TS 24.301 V11.14.0 (2015-06)"
_SPEC_TITLE_PAT = re.compile(
    r"3GPP\s+TS\s+([\d.]+)\s+V([\d.]+)\s*\((\d{4}-\d{2})\)"
)

# Version from filename: "24301-be0" → we already know this from download
_VERSION_FROM_NAME = re.compile(r"\d+-([a-z0-9]{3})")


class SpecParser:
    """Parse a 3GPP DOCX specification into a SpecDocument."""

    def parse(self, docx_path: Path) -> SpecDocument:
        """Parse a DOCX file into a SpecDocument with sections."""
        doc = docx.Document(str(docx_path))

        spec_doc = SpecDocument(source_file=str(docx_path))

        # Extract metadata from front matter
        self._extract_metadata(doc, spec_doc)

        # Parse sections from headings
        sections = self._parse_sections(doc)
        spec_doc.sections = sections

        logger.info(
            f"Parsed TS {spec_doc.spec_number} v{spec_doc.version}: "
            f"{len(sections)} sections"
        )
        return spec_doc

    def _extract_metadata(
        self, doc: docx.Document, spec_doc: SpecDocument
    ) -> None:
        """Extract spec number, version, title from front matter."""
        for p in doc.paragraphs[:80]:
            text = p.text.strip()
            if not text:
                continue

            # Look for "3GPP TS 24.301 V11.14.0 (2015-06)"
            m = _SPEC_TITLE_PAT.search(text)
            if m:
                spec_doc.spec_number = m.group(1)
                spec_doc.version = m.group(2)
                # Derive release from major version
                major = int(spec_doc.version.split(".")[0])
                spec_doc.release_num = major
                spec_doc.release = f"Release {major}"
                continue

            # Look for spec title (usually on its own line)
            if text.startswith("Non-Access") or "NAS protocol" in text:
                if not spec_doc.title:
                    spec_doc.title = text

        # Fallback: extract from filename (reliable since we control downloads)
        if not spec_doc.spec_number or not spec_doc.version:
            self._extract_from_filename(spec_doc)

    @staticmethod
    def _extract_from_filename(spec_doc: SpecDocument) -> None:
        """Extract metadata from the 3GPP filename convention.

        Filename: "24301-be0.docx" → spec=24.301, version=11.14.0
        """
        from src.standards.spec_resolver import code_to_version

        name = Path(spec_doc.source_file).stem
        # Match "24301-be0" pattern
        m = re.match(r"(\d{5,})-([a-z0-9]{3})$", name)
        if not m:
            return

        raw_num = m.group(1)
        version_code = m.group(2)

        if not spec_doc.spec_number:
            # "24301" → "24.301"
            if len(raw_num) == 5:
                spec_doc.spec_number = f"{raw_num[:2]}.{raw_num[2:]}"
            elif len(raw_num) == 6:
                spec_doc.spec_number = f"{raw_num[:2]}.{raw_num[2:]}"

        if not spec_doc.version:
            spec_doc.version = code_to_version(version_code)
            if spec_doc.version:
                major = int(spec_doc.version.split(".")[0])
                spec_doc.release_num = major
                spec_doc.release = f"Release {major}"

    def _parse_sections(self, doc: docx.Document) -> list[SpecSection]:
        """Parse all heading paragraphs into SpecSection objects."""
        sections: list[SpecSection] = []
        section_map: dict[str, SpecSection] = {}

        # First pass: collect all headings with their paragraph indices
        heading_indices: list[tuple[int, str, str]] = []
        for i, p in enumerate(doc.paragraphs):
            if not p.style.name.startswith("Heading"):
                continue
            text = p.text.strip()
            if not text:
                continue

            sec_num, title = self._parse_heading_text(text)
            if sec_num or title:
                heading_indices.append((i, sec_num, title))

        # Second pass: collect body text between headings
        for idx, (para_idx, sec_num, title) in enumerate(heading_indices):
            # Find the next heading's paragraph index
            if idx + 1 < len(heading_indices):
                next_para_idx = heading_indices[idx + 1][0]
            else:
                next_para_idx = len(doc.paragraphs)

            # Collect body text between this heading and the next
            body_parts = []
            for j in range(para_idx + 1, next_para_idx):
                p = doc.paragraphs[j]
                text = p.text.strip()
                if text:
                    body_parts.append(text)

            body = "\n".join(body_parts)

            # Determine depth from section number
            if sec_num:
                depth = sec_num.count(".") + 1
            else:
                depth = 0

            # Determine parent
            parent = self._find_parent(sec_num) if sec_num else ""

            section = SpecSection(
                number=sec_num,
                title=title,
                depth=depth,
                text=body,
                parent_number=parent,
            )
            sections.append(section)
            if sec_num:
                section_map[sec_num] = section

        # Third pass: populate children
        for sec in sections:
            if not sec.number:
                continue
            parent = section_map.get(sec.parent_number)
            if parent:
                parent.children.append(sec.number)

        return sections

    @staticmethod
    def _parse_heading_text(text: str) -> tuple[str, str]:
        """Parse a heading into (section_number, title).

        Returns ("", title) for unnumbered headings.
        """
        # Try numbered section: "5.5.1\tAttach procedure"
        m = _SECTION_NUM.match(text)
        if m:
            return m.group(1), m.group(2).strip()

        # Try annex: "Annex A (informative): GPRS timer information"
        m = _ANNEX_PAT.match(text)
        if m:
            return m.group(1), m.group(2).strip()

        # Unnumbered heading (e.g., "Foreword")
        return "", text.strip()

    @staticmethod
    def _find_parent(section_number: str) -> str:
        """Derive the parent section number.

        '5.5.1.2.5' → '5.5.1.2'
        '5' → ''
        'Annex A' → ''
        '4.2A' → '4.2'  (strip trailing letter)
        """
        if not section_number:
            return ""
        if section_number.lower().startswith("annex"):
            return ""

        # Strip trailing letter (e.g., "4.2A" → "4.2")
        cleaned = section_number.rstrip("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz")

        parts = cleaned.rsplit(".", 1)
        if len(parts) > 1:
            return parts[0]
        return ""
